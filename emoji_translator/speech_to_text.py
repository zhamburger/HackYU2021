import speech_recognition as sr
from deep_translator import GoogleTranslator
import time
import datetime
# import emoji_translate as emo

import docx
import emoji as e

e.main()
d = e.dict
# doc = docx.Document("/worddoc.docx")

# from emoji_translate.emoji_translate import Translator

# emo = Translator(exact_match_only=False, randomize=True)

def translate(lang, file):
    dest_lang = lang
    file_name = file
    # dest_lang = "english"
    # file_name = "file.docx"
    to_translate = 'I want to translate this text'
    current = 45
    start_time = datetime.datetime(100, 1, 1, 0, 0, 0)
    max_time = datetime.datetime(100, 1, 1, 0, 0, current)
    start = time.time()

    mydoc = docx.Document()
    recognizer = sr.Recognizer()
    # for i in range(2):
    x = 1
    bool = True
    while bool:
        while time.time() - start <= x:
            with sr.Microphone() as source:
                print("Talk")
                # audio_text = recognizer.record(source, duration=3)
                audio_text = recognizer.listen(source)
                try:
                    st = recognizer.recognize_google(audio_text)
                    array = st.split()
                    for i in range(len(array)):
                        if array[i] in d.keys():
                            array[i] = d[array[i]]
                    s = " ".join(array)
                    print(type(s))
                except Exception as e:
                    print("Error: " + str(e))

                print("Time over, thanks")

                try:
                    translated = GoogleTranslator(source='auto', target=dest_lang).translate(s)
                    return translated
                    # mydoc.add_paragraph(translated)
                    # mydoc.save(file_name)
                    # print(emo.emojify(text))
                except:
                    print("Sorry, I did not get that")
        x = x + 1
        if(x >= current):
            bool = False